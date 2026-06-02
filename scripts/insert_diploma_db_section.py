# -*- coding: utf-8 -*-
"""Вставка подраздела о БД (текст, таблицы, ER-рисунок) перед п. 2.4 в filippovdiplom11-1.docx."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Cm
from docx.text.paragraph import Paragraph


def bump_figure_captions_by_one(doc: Document) -> None:
    """Увеличить номер у каждой подписи «Рисунок N.» на 1 (сверху вниз по N, чтобы не затереть номера)."""
    rx = re.compile(r"^(Рисунок\s+)(\d+)(\.\s+)")
    numbered: list[tuple[int, object]] = []
    for p in doc.paragraphs:
        if not p.text:
            continue
        t = p.text.strip()
        m = rx.match(t)
        if m:
            numbered.append((int(m.group(2)), p))
    for n, p in sorted(numbered, key=lambda x: -x[0]):
        t = p.text.strip()
        p.text = rx.sub(rf"\g<1>{n + 1}\g<3>", t, count=1)


def add_table_after_paragraph(paragraph, rows_data: list[list[str]], cols: int) -> None:
    doc = paragraph.part.document
    table = doc.add_table(rows=len(rows_data), cols=cols)
    table.style = "Table Grid"
    for r, row_cells in enumerate(rows_data):
        for c, val in enumerate(row_cells):
            table.rows[r].cells[c].text = val
    tbl_el = table._tbl
    tbl_el.getparent().remove(tbl_el)
    paragraph._element.addnext(tbl_el)


def main() -> int:
    docx_path = Path(r"c:\Users\andre\Downloads\Telegram Desktop\filippovdiplom11-1.docx")
    img_path = Path(r"c:\Users\andre\OneDrive\Рабочий стол\КП\2026\БыковоГаз\docs\plan-db-schema.png")

    if not docx_path.is_file():
        print("Не найден файл:", docx_path, file=sys.stderr)
        return 1
    if not img_path.is_file():
        print("Не найден PNG схемы:", img_path, file=sys.stderr)
        return 1

    doc = Document(str(docx_path))
    anchor_2_4 = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("2.4.") and "прототипа" in p.text.lower():
            anchor_2_4 = p
            break
    if anchor_2_4 is None:
        print("Не найден абзац 2.4", file=sys.stderr)
        return 1

    marker = "\t2.3.1 Проектирование реляционной базы данных"

    if any(marker.strip() in (p.text or "") for p in doc.paragraphs):
        print("Фрагмент уже присутствует, повторная вставка отменена:", docx_path)
        return 0

    bump_figure_captions_by_one(doc)

    tables_spec: list[tuple[str, list[list[str]]]] = [
        (
            "Таблица 1 — структура отношения users",
            [
                ["№ п/п", "Поле", "Тип данных", "Назначение"],
                ["1", "user_id", "целое, PK", "Идентификатор учётной записи"],
                ["2", "login", "строка", "Логин для входа в панель администратора"],
                ["3", "password_hash", "строка", "Хэш пароля"],
                ["4", "email", "строка", "Контактный адрес электронной почты"],
                ["5", "created_at", "дата и время", "Момент создания записи"],
                ["6", "updated_at", "дата и время", "Момент последнего изменения записи"],
            ],
        ),
        (
            "Таблица 2 — структура отношения roles",
            [
                ["№ п/п", "Поле", "Тип данных", "Назначение"],
                ["1", "role_id", "целое, PK", "Идентификатор роли"],
                ["2", "name", "строка", "Наименование роли (администратор, редактор и т. п.)"],
            ],
        ),
        (
            "Таблица 3 — структура отношения user_roles",
            [
                ["№ п/п", "Поле", "Тип данных", "Назначение"],
                ["1", "user_id", "целое, FK", "Ссылка на users.user_id; часть составного PK"],
                ["2", "role_id", "целое, FK", "Ссылка на roles.role_id; часть составного PK"],
            ],
        ),
        (
            "Таблица 4 — структура отношения departments",
            [
                ["№ п/п", "Поле", "Тип данных", "Назначение"],
                ["1", "department_id", "целое, PK", "Идентификатор подразделения"],
                ["2", "name", "строка", "Наименование подразделения"],
            ],
        ),
        (
            "Таблица 5 — структура отношения team_members",
            [
                ["№ п/п", "Поле", "Тип данных", "Назначение"],
                ["1", "member_id", "целое, PK", "Идентификатор сотрудника"],
                ["2", "first_name", "строка", "Имя"],
                ["3", "last_name", "строка", "Фамилия"],
                ["4", "position", "строка", "Должность"],
                ["5", "photo", "строка / путь к файлу", "Файл фотографии"],
                ["6", "bio", "текст", "Краткая биография"],
                ["7", "department_id", "целое, FK, необязательно", "Ссылка на departments.department_id"],
                ["8", "created_at", "дата и время", "Момент добавления карточки"],
            ],
        ),
        (
            "Таблица 6 — структура отношения social_links",
            [
                ["№ п/п", "Поле", "Тип данных", "Назначение"],
                ["1", "social_id", "целое, PK", "Идентификатор типа соцсети"],
                ["2", "name", "строка", "Условное имя (например, VK, Telegram)"],
                ["3", "base_url", "строка, необязательно", "Базовый URL сервиса"],
            ],
        ),
        (
            "Таблица 7 — структура отношения member_social_links",
            [
                ["№ п/п", "Поле", "Тип данных", "Назначение"],
                ["1", "member_id", "целое, FK", "Ссылка на team_members.member_id; часть составного PK"],
                ["2", "social_id", "целое, FK", "Ссылка на social_links.social_id; часть составного PK"],
                ["3", "url_or_value", "строка", "Профиль или идентификатор в выбранной соцсети"],
            ],
        ),
        (
            "Таблица 8 — структура отношения candidate_statuses",
            [
                ["№ п/п", "Поле", "Тип данных", "Назначение"],
                ["1", "status_id", "целое, PK", "Идентификатор статуса отклика"],
                ["2", "name", "строка", "Наименование статуса (новый, в работе, отклонён и т. д.)"],
            ],
        ),
        (
            "Таблица 9 — структура отношения candidate_profiles",
            [
                ["№ п/п", "Поле", "Тип данных", "Назначение"],
                ["1", "candidate_id", "целое, PK", "Идентификатор отклика"],
                ["2", "full_name", "строка", "ФИО кандидата"],
                ["3", "phone", "строка", "Телефон"],
                ["4", "email", "строка", "Адрес электронной почты"],
                ["5", "desired_position", "строка", "Желаемая должность"],
                ["6", "message", "текст", "Сопроводительное сообщение"],
                ["7", "status_id", "целое, FK", "Ссылка на candidate_statuses.status_id"],
                ["8", "created_at", "дата и время", "Момент отправки отклика"],
            ],
        ),
        (
            "Таблица 10 — структура отношения candidate_attachments",
            [
                ["№ п/п", "Поле", "Тип данных", "Назначение"],
                ["1", "attachment_id", "целое, PK", "Идентификатор вложения"],
                ["2", "candidate_id", "целое, FK", "Ссылка на candidate_profiles.candidate_id"],
                ["3", "file_path", "строка", "Путь к сохранённому файлу резюме"],
                ["4", "uploaded_at", "дата и время", "Момент загрузки файла"],
            ],
        ),
        (
            "Таблица 11 — структура отношения contact_messages",
            [
                ["№ п/п", "Поле", "Тип данных", "Назначение"],
                ["1", "message_id", "целое, PK", "Идентификатор обращения"],
                ["2", "name", "строка", "Имя отправителя"],
                ["3", "phone_or_email", "строка", "Телефон и/или электронная почта"],
                ["4", "subject", "строка, необязательно", "Тема сообщения"],
                ["5", "message", "текст", "Текст обращения"],
                ["6", "created_at", "дата и время", "Момент отправки формы"],
                [
                    "7",
                    "processed_by_user_id",
                    "целое, FK, необязательно",
                    "Пользователь (users), зафиксировавший обработку обращения",
                ],
            ],
        ),
        (
            "Таблица 12 — структура отношения contact_message_attachments",
            [
                ["№ п/п", "Поле", "Тип данных", "Назначение"],
                ["1", "attachment_id", "целое, PK", "Идентификатор вложения"],
                ["2", "message_id", "целое, FK", "Ссылка на contact_messages.message_id"],
                ["3", "file_path", "строка", "Путь к прикреплённому файлу"],
                ["4", "uploaded_at", "дата и время", "Момент загрузки"],
            ],
        ),
    ]

    head = anchor_2_4.insert_paragraph_before(
        "Предусмотренная логическая модель согласована с разделами сайта, где требуется хранение "
        "структурированных данных: панель администратора, раздел «Состав команды», приём откликов "
        "на вакансии и форма обратной связи. В реализованном прототипе используется встроенная СУБД "
        "SQLite (файл базы данных в каталоге проекта Django), для промышленного развёртывания без "
        "изменения кода приложения допускается замена движка на PostgreSQL средствами настроек Django."
    )

    for caption, rows in reversed(tables_spec):
        cap_p = head.insert_paragraph_before(caption)
        add_table_after_paragraph(cap_p, rows, 4)
        head = cap_p

    head = head.insert_paragraph_before(
        "Структура полей поимённых отношений приведена в таблицах 1–12 (имена отношений приведены "
        "в инфологической нотации; типы данных указаны в логическом виде, конкретные типы столбцов "
        "в СУБД определяются при физической реализации)."
    )

    fig_cap = head.insert_paragraph_before(
        "Рисунок 1. Логическая ER-схема базы данных (сотрудники, кандидаты, обращения, ролевая модель)"
    )

    pic_p = fig_cap.insert_paragraph_before("")
    pic_p.paragraph_format.space_after = None
    run = pic_p.add_run()
    run.add_picture(str(img_path), width=Cm(16.0))

    after_cap_el = OxmlElement("w:p")
    fig_cap._element.addnext(after_cap_el)
    ref_after = Paragraph(after_cap_el, fig_cap._parent)
    ref_after.add_run(
        "На рисунке 1 представлена логическая ER-схема: связи «один-ко-многим» между подразделениями "
        "и карточками сотрудников, между статусом и откликами кандидатов, между обращением и файлами; "
        "связи «многие-ко-многим» пользователей и ролей, сотрудников и соцсетей развёрнуты через "
        "связующие отношения user_roles и member_social_links."
    )

    head = pic_p.insert_paragraph_before(
        "Логическая ER-диаграмма построена в редакторе diagrams.net (draw.io); на схеме показаны "
        "имена отношений, ключи и основные атрибуты."
    )

    head = head.insert_paragraph_before(marker)

    doc.save(str(docx_path))
    print("Обновлён:", docx_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
