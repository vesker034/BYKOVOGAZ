# -*- coding: utf-8 -*-
"""Генерация «Приложение 3. Листинг программного кода» для проекта БыковоГаз."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
REF_DOC = Path(r"c:\Users\andre\Downloads\Telegram Desktop\Приложение_З_Лышков М.И.docx")
OUTPUT = ROOT / "docs" / "Приложение_3_БыковоГаз.docx"

APPENDIX_FILES: list[tuple[str, Path]] = [
    ("views.py", ROOT / "main" / "views.py"),
    ("config/urls.py", ROOT / "config" / "urls.py"),
    ("main/urls.py", ROOT / "main" / "urls.py"),
    ("models.py", ROOT / "main" / "models.py"),
    ("admin.py", ROOT / "main" / "admin.py"),
    ("forms.py", ROOT / "main" / "forms.py"),
    ("cms_seed.py", ROOT / "main" / "cms_seed.py"),
    ("style.css", ROOT / "static" / "css" / "style.css"),
    ("figma-pages.css", ROOT / "static" / "css" / "figma-pages.css"),
    ("main.js", ROOT / "static" / "js" / "main.js"),
    ("about_company.html", ROOT / "templates" / "pages" / "about_company.html"),
    ("works.html", ROOT / "templates" / "pages" / "works.html"),
    ("home.html", ROOT / "templates" / "pages" / "home.html"),
    ("career.html", ROOT / "templates" / "pages" / "career.html"),
    ("contacts.html", ROOT / "templates" / "pages" / "contacts.html"),
]

FONT_NAME = "Times New Roman"
TITLE_SIZE = Pt(16)
HEADING_SIZE = Pt(14)
CODE_SIZE = Pt(12)


def _apply_page_setup(doc: Document) -> None:
    if REF_DOC.exists():
        ref = Document(str(REF_DOC))
        ref_sec = ref.sections[0]
        for section in doc.sections:
            section.page_width = ref_sec.page_width
            section.page_height = ref_sec.page_height
            section.left_margin = ref_sec.left_margin
            section.right_margin = ref_sec.right_margin
            section.top_margin = ref_sec.top_margin
            section.bottom_margin = ref_sec.bottom_margin
    else:
        section = doc.sections[0]
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)


def _set_run_font(run, *, size: Pt, bold: bool | None = None) -> None:
    run.font.name = FONT_NAME
    run.font.size = size
    if bold is not None:
        run.bold = bold


def _add_title(doc: Document, text: str, *, align_right: bool = False, indent: bool = False) -> None:
    p = doc.add_paragraph()
    if align_right:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    _set_run_font(run, size=TITLE_SIZE if align_right else HEADING_SIZE, bold=True)


def _add_file_heading(doc: Document, filename: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(filename)
    _set_run_font(run, size=HEADING_SIZE, bold=True)


def _add_code_line(doc: Document, line: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(line)
    _set_run_font(run, size=CODE_SIZE)


def _read_source(path: Path) -> list[str]:
    text = path.read_text(encoding="utf-8")
    return text.splitlines()


def build() -> Path:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _apply_page_setup(doc)

    _add_title(doc, "Приложение 3", align_right=True)
    _add_title(doc, "Листинг программного кода", indent=True)

    missing: list[str] = []
    for display_name, source_path in APPENDIX_FILES:
        if not source_path.is_file():
            missing.append(str(source_path))
            continue
        _add_file_heading(doc, display_name)
        for line in _read_source(source_path):
            _add_code_line(doc, line)

    if missing:
        raise FileNotFoundError("Не найдены файлы:\n" + "\n".join(missing))

    doc.save(str(OUTPUT))
    return OUTPUT


if __name__ == "__main__":
    out = build()
    lines = sum(len(_read_source(p)) for _, p in APPENDIX_FILES if p.is_file())
    print(f"Сохранено: {out}")
    print(f"Файлов: {len(APPENDIX_FILES)}, строк кода: {lines}")
