# -*- coding: utf-8 -*-
"""Восстановление порядка абзацев 2.9 / 2.10 и иллюстраций (устойчиво к сбитому порядку)."""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

DOCX_PATH = Path(r"c:\Users\andre\Downloads\Telegram Desktop\filippovdiplom11-1.docx")


def _has_drawing(p: Paragraph) -> bool:
    for run in p.runs:
        if run._element.findall(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
        ):
            return True
    return False


def _find_idx(doc: Document, pred, lo: int = 0, hi: int | None = None) -> int:
    paras = doc.paragraphs
    hi = hi if hi is not None else len(paras)
    for i in range(lo, hi):
        if pred(paras[i]):
            return i
    raise LookupError("абзац не найден")


def _cluster_by_gaps(indices: list[int], gap: int = 3) -> list[list[int]]:
    if not indices:
        return []
    s = sorted(indices)
    out: list[list[int]] = [[s[0]]]
    for x in s[1:]:
        if x - out[-1][-1] <= gap:
            out[-1].append(x)
        else:
            out.append([x])
    return out


def main() -> int:
    if not DOCX_PATH.is_file():
        print("Не найден:", DOCX_PATH, file=sys.stderr)
        return 1

    doc = Document(str(DOCX_PATH))
    paras = doc.paragraphs

    i32 = _find_idx(doc, lambda p: (p.text or "").strip().startswith("Рисунок 32."))
    i11 = _find_idx(doc, lambda p: (p.text or "").strip().startswith("2.11 Развёртывание"))

    chunk_lo = i32 + 1
    chunk_hi = i11
    chunk_indices = list(range(chunk_lo, chunk_hi))

    i_h9 = _find_idx(
        doc,
        lambda p: (p.text or "").strip() == "2.9 Тестирование веб-сайта",
        chunk_lo,
        chunk_hi,
    )
    i_h10 = _find_idx(
        doc,
        lambda p: (p.text or "").strip() == "2.10 Оптимизация веб-сайта",
        chunk_lo,
        chunk_hi,
    )
    i_t1 = _find_idx(
        doc,
        lambda p: (p.text or "").strip().startswith("Тестирование проводилось вручную"),
        chunk_lo,
        chunk_hi,
    )
    i_t2 = _find_idx(
        doc,
        lambda p: (p.text or "").strip().startswith("Кросс-браузерная проверка выполнена"),
        chunk_lo,
        chunk_hi,
    )
    i_o1 = _find_idx(
        doc,
        lambda p: (p.text or "").strip().startswith("Оптимизация выполнялась по нескольким направлениям"),
        chunk_lo,
        chunk_hi,
    )
    i_o2 = _find_idx(
        doc,
        lambda p: (p.text or "").strip().startswith("На стороне сервера включена раздача статических файлов"),
        chunk_lo,
        chunk_hi,
    )
    i_o3 = _find_idx(
        doc,
        lambda p: (p.text or "").strip().startswith("Рекомендации для промышленной эксплуатации"),
        chunk_lo,
        chunk_hi,
    )
    i_c33 = _find_idx(doc, lambda p: (p.text or "").strip().startswith("Рисунок 33."), chunk_lo, chunk_hi)
    i_c34 = _find_idx(doc, lambda p: (p.text or "").strip().startswith("Рисунок 34."), chunk_lo, chunk_hi)
    i_c35 = _find_idx(doc, lambda p: (p.text or "").strip().startswith("Рисунок 35."), chunk_lo, chunk_hi)
    i_c36 = _find_idx(doc, lambda p: (p.text or "").strip().startswith("Рисунок 36."), chunk_lo, chunk_hi)

    draw_only = [
        i
        for i in chunk_indices
        if _has_drawing(paras[i]) and not (paras[i].text or "").strip()
    ]
    pre_draws = sorted(i for i in draw_only if i < i_h10)
    post_draws = sorted(i for i in draw_only if i > i_h10)

    cl_pre = _cluster_by_gaps(pre_draws, gap=2)
    cl_post = _cluster_by_gaps(post_draws, gap=2)

    lone = [c for c in cl_pre if len(c) == 1]
    multi = [c for c in cl_pre if len(c) > 1]
    if len(lone) != 1 or len(multi) != 1:
        print(
            "Неожиданный набор иллюстраций в блоке тестирования (ожидались 1 одиночная и 1 пара).",
            "кластеры:",
            cl_pre,
            file=sys.stderr,
        )
        return 1
    img_form = lone[0][0]
    imgs_news = sorted(multi[0])

    if len(cl_post) < 2:
        print("Неожиданный набор иллюстраций в блоке оптимизации:", cl_post, file=sys.stderr)
        return 1
    cl_post_s = sorted(cl_post, key=lambda c: min(c))
    img_net = cl_post_s[0][0]
    img_lh = cl_post_s[1][0]

    empties = [i for i in chunk_indices if not (paras[i].text or "").strip() and not _has_drawing(paras[i])]

    def first_empty_between(a: int, b: int) -> list[int]:
        lo, hi = sorted((a, b))
        return [i for i in empties if lo < i < hi][:1]

    e1 = first_empty_between(i_t2, img_form)
    e2 = first_empty_between(i_c33, min(imgs_news))
    e3 = first_empty_between(i_o3, img_net)
    e4 = first_empty_between(i_c35, img_lh)

    ordered: list[int] = [
        i_h9,
        i_t1,
        i_t2,
        *e1,
        img_form,
        i_c33,
        *e2,
        *imgs_news,
        i_c34,
        i_h10,
        i_o1,
        i_o2,
        i_o3,
        *e3,
        img_net,
        i_c35,
        *e4,
        img_lh,
        i_c36,
    ]

    ordered_unique = list(dict.fromkeys(ordered))
    if len(ordered_unique) != len(chunk_indices):
        missing = sorted(set(chunk_indices) - set(ordered_unique))
        extra = sorted(set(ordered_unique) - set(chunk_indices))
        print("Несовпадение числа абзацев. missing=", missing, "extra=", extra, file=sys.stderr)
        if missing:
            ordered_unique.extend(missing)

    anchor = paras[i32]._element
    parent = anchor.getparent()
    if parent is None:
        return 1

    elements = [paras[i]._element for i in ordered_unique]
    for el in elements:
        p = el.getparent()
        if p is not None:
            p.remove(el)

    cur = anchor
    for el in elements:
        cur.addnext(el)
        cur = el

    doc.save(str(DOCX_PATH))
    print("Порядок восстановлен:", DOCX_PATH)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
