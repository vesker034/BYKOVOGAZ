# -*- coding: utf-8 -*-
"""Подмена подписи «Рисунок 42.» (физическая структура проекта) в filippovdiplom11-1.docx на docs/site-structure-physical.png."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docx.text.paragraph import Paragraph

DOCX_PATH = Path(r"c:\Users\andre\Downloads\Telegram Desktop\filippovdiplom11-1.docx")
IMG_PATH = Path(__file__).resolve().parents[1] / "docs" / "site-structure-physical.png"
CAPTION_RE = re.compile(r"^Рисунок\s*42\.\s")

IMG_WIDTH_CM = 16.0


def _clear_paragraph(paragraph: Paragraph) -> None:
    el = paragraph._element
    for child in list(el):
        el.remove(child)


def _paragraph_has_drawing(paragraph: Paragraph) -> bool:
    for run in paragraph.runs:
        if run._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"):
            return True
    return False


def main() -> int:
    if not DOCX_PATH.is_file():
        print("Не найден файл:", DOCX_PATH, file=sys.stderr)
        return 1
    if not IMG_PATH.is_file():
        print("Не найден PNG:", IMG_PATH, file=sys.stderr)
        return 1

    doc = Document(str(DOCX_PATH))
    cap_idx: int | None = None
    for i, p in enumerate(doc.paragraphs):
        if CAPTION_RE.match((p.text or "").strip()):
            cap_idx = i
            break
    if cap_idx is None:
        print("Не найдена подпись «Рисунок 42.»", file=sys.stderr)
        return 1

    if cap_idx < 1:
        print("Некорректная позиция подписи к рисунку", file=sys.stderr)
        return 1

    img_p = doc.paragraphs[cap_idx - 1]
    if not _paragraph_has_drawing(img_p) and (img_p.text or "").strip():
        print(
            "Абзац перед подписью не похож на блок с иллюстрацией; остановка без изменений.",
            file=sys.stderr,
        )
        return 1

    _clear_paragraph(img_p)
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_p.add_run()
    run.add_picture(str(IMG_PATH), width=Cm(IMG_WIDTH_CM))

    doc.save(str(DOCX_PATH))
    print("Обновлён рисунок в:", DOCX_PATH)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
