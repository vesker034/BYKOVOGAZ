# -*- coding: utf-8 -*-
"""П. 1.4 (обзор технологий): вводный абзац с актуальным стеком + правка версии Django."""

from pathlib import Path

from docx import Document

DOC_PATH = Path(r"c:\Users\andre\Downloads\Telegram Desktop\filippovdiplom11-1.docx")
TITLE = "Обзор технологий и средств разработки веб-сайта"
STACK_INTRO = (
    "Для реализации корпоративного сайта выбран следующий технологический стек: язык Python 3; "
    "веб-фреймворк Django 6.0.3 (маршрутизация запросов, ORM, шаблоны, встроенная "
    "административная панель, интернационализация интерфейса); СУБД SQLite; клиентская часть "
    "на HTML5, CSS3 и JavaScript с применением Bootstrap 5; статические таблицы стилей и "
    "клиентские сценарии проекта; для опытной эксплуатации на сервере — связка Gunicorn и "
    "обратного прокси Nginx; контроль версий Git и публикация исходного кода на GitHub; "
    "проектирование макетов в Figma; разработка в Visual Studio Code."
)
REF_BODY_INDEX = 255


def _copy_paragraph_format(target, source):
    dst = target.paragraph_format
    src = source.paragraph_format
    dst.left_indent = src.left_indent
    dst.right_indent = src.right_indent
    dst.first_line_indent = src.first_line_indent
    dst.space_before = src.space_before
    dst.space_after = src.space_after
    dst.line_spacing = src.line_spacing
    dst.line_spacing_rule = src.line_spacing_rule
    dst.keep_together = src.keep_together
    dst.keep_with_next = src.keep_with_next
    dst.page_break_before = src.page_break_before
    dst.widow_control = src.widow_control
    target.alignment = source.alignment


def _apply_runs_like_ref(paragraph, ref_paragraph):
    paragraph.style = ref_paragraph.style
    ref_runs = ref_paragraph.runs
    if not ref_runs:
        return
    ref = ref_runs[0]
    for run in paragraph.runs:
        if ref.font.name:
            run.font.name = ref.font.name
        if ref.font.size:
            run.font.size = ref.font.size
        if ref.bold is not None:
            run.bold = ref.bold
        if ref.italic is not None:
            run.italic = ref.italic


def main():
    doc = Document(str(DOC_PATH))
    ref_body = doc.paragraphs[REF_BODY_INDEX]

    title_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() != TITLE:
            continue
        if i + 1 >= len(doc.paragraphs):
            continue
        nxt = doc.paragraphs[i + 1].text.lstrip()
        if nxt.startswith("HTML5") or nxt.startswith("Для реализации корпоративного сайта выбран"):
            title_idx = i
            break
    if title_idx is None:
        raise SystemExit(f"Не найден заголовок раздела в тексте: {TITLE!r}")

    next_p = doc.paragraphs[title_idx + 1]
    if next_p.text.strip().startswith("Для реализации корпоративного сайта выбран следующий"):
        print("Вводный абзац о стеке уже присутствует после заголовка 1.4.")
    else:
        new_p = next_p.insert_paragraph_before(STACK_INTRO)
        _copy_paragraph_format(new_p, ref_body)
        _apply_runs_like_ref(new_p, ref_body)

    replaced = 0
    for p in doc.paragraphs:
        if "версии 6.x" in p.text or "Django 6.x" in p.text:
            old = p.text
            new = old.replace("версии 6.x", "6.0.3").replace("Django 6.x", "Django 6.0.3")
            if new != old:
                p.text = new
                replaced += 1
        if "Django (6.0.3):" in p.text:
            p.text = p.text.replace("Django (6.0.3):", "Django 6.0.3:")

    doc.save(str(DOC_PATH))
    print(f"Сохранено: {DOC_PATH}; замен по «6.x»: {replaced}")


if __name__ == "__main__":
    main()
