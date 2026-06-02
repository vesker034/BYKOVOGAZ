# -*- coding: utf-8 -*-
"""Добавление в п. 2.9 текста и двух рисунков про мобильную адаптацию; сдвиг нумерации рис. 35 и выше на +2."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document

DOCX_PATH = Path(r"c:\Users\andre\Downloads\Telegram Desktop\filippovdiplom11-1.docx")

MARKER = "адаптивной вёрстки: сравнивались компоновка страниц"

CAP_RX = re.compile(r"^(Рисунок\s+)(\d+)(\..*)$", re.DOTALL)

MOBILE_BODY = (
    "Отдельно фиксировалась работа адаптивной вёрстки: сравнивались компоновка страниц при ширине окна, "
    "характерной для настольного монитора (не менее 1280 px), и при эмуляции смартфона в инструментах "
    "разработчика (ширина области просмотра ориентировочно 360–375 px). На рисунках ниже приведены примеры "
    "отображения разделов сайта при мобильной ширине."
)

PH1 = (
    "[Вставить иллюстрацию: главная страница в режиме эмуляции мобильного устройства в инструментах "
    "разработчика браузера (ширина области просмотра около 375 px). Формат PNG или JPG, ширина на странице "
    "14–16 см, подпись оформить как у соседних рисунков.]"
)

PH2 = (
    "[Вставить иллюстрацию: другой раздел сайта (например, «Контакты», «Новости» или шапка с мобильным меню) "
    "при той же узкой ширине; должно быть видно отличие от настольной компоновки. Формат PNG или JPG.]"
)

CAP35 = (
    "Рисунок 35. Главная страница при эмуляции мобильной ширины окна браузера (~375 px)"
)

CAP36 = (
    "Рисунок 36. Раздел сайта при мобильной ширине окна (адаптивная компоновка)"
)


def _ref_testing_body(doc: Document):
    for p in doc.paragraphs:
        if (p.text or "").strip().startswith("Тестирование проводилось вручную"):
            return p
    raise LookupError("эталонный абзац тестирования")


def _ref_placeholder(doc: Document):
    for p in doc.paragraphs:
        if "[Вставить иллюстрацию:" in (p.text or ""):
            return p
    raise LookupError("эталон плейсхолдера")


def _ref_caption(doc: Document):
    for p in doc.paragraphs:
        if (p.text or "").strip().startswith("Рисунок 33."):
            return p
    raise LookupError("эталон подписи 33")


def _copy_para_style(dst, src) -> None:
    dst.style = src.style
    dst.paragraph_format.left_indent = src.paragraph_format.left_indent
    dst.paragraph_format.right_indent = src.paragraph_format.right_indent
    dst.paragraph_format.first_line_indent = src.paragraph_format.first_line_indent
    dst.paragraph_format.space_before = src.paragraph_format.space_before
    dst.paragraph_format.space_after = src.paragraph_format.space_after
    dst.paragraph_format.line_spacing = src.paragraph_format.line_spacing
    dst.paragraph_format.line_spacing_rule = src.paragraph_format.line_spacing_rule
    dst.alignment = src.alignment


def _apply_run_like(src_para, dst_para, *, italic: bool | None = None, bold: bool | None = None) -> None:
    if not src_para.runs:
        return
    ref = src_para.runs[0]
    for run in dst_para.runs:
        if ref.font.name:
            run.font.name = ref.font.name
        if ref.font.size:
            run.font.size = ref.font.size
        if bold is not None:
            run.bold = bold
        elif ref.bold is not None:
            run.bold = ref.bold
        if italic is not None:
            run.italic = italic
        elif ref.italic is not None:
            run.italic = ref.italic


def _bump_captions_from(doc: Document, from_n: int, delta: int) -> None:
    pending: list[tuple[int, object, re.Match[str]]] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        m = CAP_RX.match(t)
        if m:
            n = int(m.group(2))
            if n >= from_n:
                pending.append((n, p, m))
    for n, p, m in sorted(pending, key=lambda x: -x[0]):
        p.text = m.group(1) + str(n + delta) + m.group(3)


def main() -> int:
    if not DOCX_PATH.is_file():
        print("Не найден:", DOCX_PATH, file=sys.stderr)
        return 1

    doc = Document(str(DOCX_PATH))
    if any(MARKER in (p.text or "") for p in doc.paragraphs):
        print("Фрагмент про мобильную адаптацию уже присутствует:", DOCX_PATH)
        return 0

    ref_body = _ref_testing_body(doc)
    ref_ph = _ref_placeholder(doc)
    ref_cap = _ref_caption(doc)

    anchor = None
    for p in doc.paragraphs:
        if (p.text or "").strip().startswith("2.10 Оптимизация веб-сайта"):
            anchor = p
            break
    if anchor is None:
        print("Не найден заголовок 2.10", file=sys.stderr)
        return 1

    _bump_captions_from(doc, 35, 2)

    stack = [
        (MOBILE_BODY, "body"),
        ("", "empty"),
        (PH1, "placeholder"),
        ("", "empty"),
        (CAP35, "caption"),
        ("", "empty"),
        (PH2, "placeholder"),
        ("", "empty"),
        (CAP36, "caption"),
    ]

    for text, kind in stack:
        np = anchor.insert_paragraph_before(text)
        if kind == "body":
            _copy_para_style(np, ref_body)
            _apply_run_like(ref_body, np)
        elif kind == "placeholder":
            _copy_para_style(np, ref_ph)
            _apply_run_like(ref_ph, np, italic=True)
        elif kind == "caption":
            _copy_para_style(np, ref_cap)
            _apply_run_like(ref_cap, np, bold=True)
        else:
            _copy_para_style(np, ref_body)
            for r in np.runs:
                r.text = ""

    doc.save(str(DOCX_PATH))
    print("Добавлены абзацы и плейсхолдеры рис. 35–36; нумерация рис. 35+ сдвинута на 2:", DOCX_PATH)
    print("Подпись диаграммы физической структуры теперь «Рисунок 42.» — при вставке PNG используйте обновлённый скрипт.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
